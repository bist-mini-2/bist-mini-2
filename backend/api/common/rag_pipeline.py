"""3개 도메인의 공통 RAG 임베딩 검색 파이프라인 및 LangChain 도구(Tools)를 관리하는 모듈입니다."""

import io
import logging
from typing import Any, Dict, List
import zipfile
import xml.etree.ElementTree as ET
import pypdfium2 as pdfium
from docx import Document as DocxDocument
from langchain.embeddings import init_embeddings
from langchain_postgres import PGVector
from langchain.tools import tool, ToolRuntime
from langchain_core.messages import ToolMessage
from langgraph.types import Command
from langchain_community.tools.tavily_search import TavilySearchResults
from api.common.config import settings

logger = logging.getLogger(__name__)

EMBED_MODEL = "openai:text-embedding-3-large"
CONNECTION = settings.PGVECTOR_URL
CHUNK_SIZE = 800
CHUNK_OVERLAP = 150
SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".csv", ".docx", ".doc"}


DOMAIN_COLLECTIONS = {
    "bio": "bio_embeddings",
    "cs": "cs_embeddings",
    "astronomy": "astronomy_embeddings"
}


class CommonRagPipeline:
    """공통 RAG 파이프라인 서비스 클래스입니다.

    3개 학술 도메인(bio, cs, astronomy)의 논문 임베딩 검색 및
    유사도 계산 처리를 단일 langchain pgvector 구조로 추상화하여 수행합니다.
    """

    def __init__(self) -> None:
        """CommonRagPipeline 인스턴스를 초기화합니다."""
        self.logger = logging.getLogger(f"{__name__}.CommonRagPipeline")
        self._embeddings = None

    def get_embeddings(self):
        """임베딩 인스턴스를 지연 로딩으로 초기화 및 공유합니다.

        Returns:
            Embeddings: 초기화된 임베딩 생성 모델 인스턴스.
        """
        if self._embeddings is None:
            self._embeddings = init_embeddings(model=EMBED_MODEL)
        return self._embeddings

    async def similarity_search(
        self,
        domain: str,
        query: str,
        k: int = 10
    ) -> List[Dict[str, Any]]:
        """지정된 도메인 컬렉션에서 질의어와 유사도가 높은 문서를 검색합니다.

        Args:
            domain (str): 학술 도메인 명 ('bio', 'cs', 'astronomy').
            query (str): 사용자의 질의 텍스트.
            k (int): 반환할 상위 문서의 개수. Defaults to 3.

        Returns:
            List[Dict[str, Any]]: 유사 논문 청크 정보 목록. 각 딕셔너리는 다음 키를 가집니다:
                - doc_id (str): 논문 고유 ID (arxiv_id).
                - title (str): 논문 제목.
                - text_chunk (str): 문서 텍스트 청크 본문.
                - score (float): 코사인 유사도 점수 (1.0 - distance).

        Raises:
            ValueError: 지원하지 않는 도메인이 입력된 경우.
        """
        if domain not in DOMAIN_COLLECTIONS:
            raise ValueError(f"지원하지 않는 도메인입니다: {domain}")

        collection_name = DOMAIN_COLLECTIONS[domain]
        self.logger.info(
            f"RAG Similarity Search: Domain='{domain}', Collection='{collection_name}', k={k}")

        vectorstore = PGVector(
            embeddings=self.get_embeddings(),
            collection_name=collection_name,
            connection=CONNECTION,
            async_mode=True,
        )

        results = await vectorstore.asimilarity_search_with_score(query, k=k)

        import math

        def safe_str(val: Any) -> str:
            if val is None:
                return ""
            if isinstance(val, float) and math.isnan(val):
                return ""
            return str(val)

        SIMILARITY_THRESHOLD = 0.4  # 이 값 미만(약한 매칭)은 관련 없음으로 보고 제외

        formatted_results = []
        for doc, score in results:
            meta = doc.metadata or {}
            arxiv_id = meta.get("arxiv_id") or meta.get("doc_id") or ""
            title = meta.get("title", "")

            # 코사인 유사도 계산 (1.0 - 거리 점수)
            similarity = round(1.0 - score, 4)

            # 임계값 미만은 제외
            if similarity < SIMILARITY_THRESHOLD:
                continue

            formatted_results.append({
                "doc_id": safe_str(arxiv_id),
                "title": safe_str(title),
                "text_chunk": safe_str(doc.page_content),
                "score": similarity
            })

        self.logger.info(f"임계값({SIMILARITY_THRESHOLD}) 통과: {len(formatted_results)}/{len(results)}건")
        return formatted_results


# 싱글톤 인스턴스 생성
common_rag_pipeline = CommonRagPipeline()


@tool
async def search_bio_papers(
    query: str,
    runtime: ToolRuntime,
    k: int = 10
) -> Command:
    """생명공학·유전체학(q-bio.GN) 논문 데이터베이스에서 관련 내용을 검색합니다.

    유전체학, 유전자 매핑, 시퀀싱 분석 등 생명공학/유전학 질문에 대답하거나 참고 자료가 필요할 때 이 툴을 사용하세요.

    Args:
        query (str): 검색할 핵심 키워드 또는 영어 학술 질문.
        runtime (ToolRuntime): LangGraph 라이브러리에서 관리하는 툴 런타임 컨텍스트.
        k (int): 반환할 상위 유사 논문 청크의 개수. 기본값은 3.

    Returns:
        Command: LangGraph의 메시지 상태(messages)와 검색 출처(sources)를 업데이트하는 흐름 제어 명령 객체.
    """
    results = await common_rag_pipeline.similarity_search("bio", query, k=k)

    if not results:
        msg = f"q-bio.GN 논문에서 '{query}'와 관련된 내용을 찾을 수 없습니다."
        return Command(update={
            "messages": [ToolMessage(content=msg, tool_call_id=runtime.tool_call_id)],
        })

    output_lines = [f"검색 결과: '{query}' (q-bio.GN 논문)\n", "=" * 80]
    sources = []
    for idx, r in enumerate(results, 1):
        arxiv_id = r["doc_id"]
        title = r["title"]
        score = r["score"]
        output_lines.append(f"\n[논문 {idx}] (유사도: {score:.4f})")
        output_lines.append(f"제목: {title}")
        output_lines.append(f"arXiv ID: {arxiv_id}")
        output_lines.append(f"\n내용:\n{r['text_chunk']}\n")
        output_lines.append("-" * 80)
        text_chunk = r.get("text_chunk")
        if not isinstance(text_chunk, str):
            text_chunk = ""
        snippet = " ".join(text_chunk.split())
        if len(snippet) > 160:
            snippet = snippet[:160].rstrip() + "…"
        sources.append({"arxiv_id": arxiv_id, "title": title, "summary": snippet})

    tool_text = "\n".join(output_lines)
    return Command(update={
        "messages": [ToolMessage(content=tool_text, tool_call_id=runtime.tool_call_id)],
        "sources": sources,
    })


@tool
async def search_cs_papers(
    query: str,
    runtime: ToolRuntime,
    k: int = 10
) -> Command:
    """컴퓨터 과학(cs.NE) 관련 학술 논문 데이터베이스에서 관련 내용을 검색합니다.

    인공신경망, 진화 알고리즘, 딥러닝 등에 대한 정보가 필요하거나 질문에 대답할 때 이 툴을 사용하세요.

    Args:
        query (str): 검색할 핵심 키워드 또는 영어 학술 질문.
        runtime (ToolRuntime): LangGraph 라이브러리에서 관리하는 툴 런타임 컨텍스트.
        k (int): 반환할 상위 유사 논문 청크의 개수. 기본값은 3.

    Returns:
        Command: LangGraph의 메시지 상태(messages)와 검색 출처(sources)를 업데이트하는 흐름 제어 명령 객체.
    """
    results = await common_rag_pipeline.similarity_search("cs", query, k=k)

    if not results:
        msg = f"cs.NE 논문에서 '{query}'와 관련된 내용을 찾을 수 없습니다."
        return Command(update={
            "messages": [ToolMessage(content=msg, tool_call_id=runtime.tool_call_id)],
        })

    output_lines = [f"검색 결과: '{query}' (cs.NE 논문)\n", "=" * 80]
    sources = []
    for idx, r in enumerate(results, 1):
        arxiv_id = r["doc_id"]
        title = r["title"]
        score = r["score"]

        output_lines.append(f"\n[논문 {idx}] (유사도: {score:.4f})")
        output_lines.append(f"제목: {title}")
        output_lines.append(f"arXiv ID: {arxiv_id}")
        output_lines.append(f"\n내용:\n{r['text_chunk']}\n")
        output_lines.append("-" * 80)

        text_chunk = r.get("text_chunk")
        if not isinstance(text_chunk, str):
            text_chunk = ""
        snippet = " ".join(text_chunk.split())
        if len(snippet) > 160:
            snippet = snippet[:160].rstrip() + "…"
        sources.append({"arxiv_id": arxiv_id, "title": title, "summary": snippet})

    tool_text = "\n".join(output_lines)
    return Command(update={
        "messages": [ToolMessage(content=tool_text, tool_call_id=runtime.tool_call_id)],
        "sources": sources,
    })


@tool
async def search_astronomy_papers(
    query: str,
    runtime: ToolRuntime,
    k: int = 10
) -> Command:
    """지구 및 행성 천체물리학(astro-ph.EP) 논문 데이터베이스에서 관련 내용을 검색합니다.

    행성 대기, Mars 지질학, 행성 형성 등에 관한 질문에 대답하거나 참고 자료가 필요할 때 이 툴을 사용하세요.

    Args:
        query (str): 검색할 핵심 키워드 또는 영어 학술 질문.
        runtime (ToolRuntime): LangGraph 라이브러리에서 관리하는 툴 런타임 컨텍스트.
        k (int): 반환할 상위 유사 논문 청크의 개수. 기본값은 3.

    Returns:
        Command: LangGraph의 메시지 상태(messages)와 검색 출처(sources)를 업데이트하는 흐름 제어 명령 객체.
    """
    results = await common_rag_pipeline.similarity_search("astronomy", query, k=k)

    if not results:
        msg = f"astro-ph.EP 논문에서 '{query}'와 관련된 내용을 찾을 수 없습니다."
        return Command(update={
            "messages": [ToolMessage(content=msg, tool_call_id=runtime.tool_call_id)],
        })

    output_lines = [f"검색 결과: '{query}' (astro-ph.EP 논문)\n", "=" * 80]
    sources = []
    for idx, r in enumerate(results, 1):
        arxiv_id = r["doc_id"]
        title = r["title"]
        score = r["score"]

        output_lines.append(f"\n[논문 {idx}] (유사도: {score:.4f})")
        output_lines.append(f"제목: {title}")
        output_lines.append(f"arXiv ID: {arxiv_id}")
        output_lines.append(f"\n내용:\n{r['text_chunk']}\n")
        output_lines.append("-" * 80)

        text_chunk = r.get("text_chunk")
        if not isinstance(text_chunk, str):
            text_chunk = ""
        snippet = " ".join(text_chunk.split())
        if len(snippet) > 160:
            snippet = snippet[:160].rstrip() + "…"
        sources.append({"arxiv_id": arxiv_id, "title": title, "summary": snippet})

    tool_text = "\n".join(output_lines)
    return Command(update={
        "messages": [ToolMessage(content=tool_text, tool_call_id=runtime.tool_call_id)],
        "sources": sources,
    })


def _collection_name(gem_id: str) -> str:
    """Gem별 pgvector 컬렉션 이름을 반환합니다.

    Args:
        gem_id (str): 대상 Gem의 고유 UUID.

    Returns:
        str: pgvector 컬렉션명.
    """
    return f"gem_{gem_id}_files"


def _extract_text(filename: str, file_bytes: bytes) -> str:
    """파일 형식에 따라 텍스트를 추출합니다.

    Args:
        filename (str): 원본 파일명 (확장자 판별에 사용).
        file_bytes (bytes): 파일 바이너리 데이터.

    Returns:
        str: 추출된 텍스트 문자열.
    """
    fname = filename.lower()

    if fname.endswith(".pdf"):
        try:
            pdf = pdfium.PdfDocument(io.BytesIO(file_bytes))
            pages = []
            for page in pdf:
                textpage = page.get_textpage()
                pages.append(textpage.get_text_range() or "")
            return "\n".join(p for p in pages if p.strip())
        except Exception as e:
            logger.error(f"PDF 텍스트 추출 실패: {e}")
            return ""

    if fname.endswith(".docx") or fname.endswith(".doc"):
        try:
            # 1차 시도: 초고속 XML 추출 기법
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as docx:
                xml_content = docx.read('word/document.xml')
            root = ET.fromstring(xml_content)
            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            texts = [node.text for node in root.findall('.//w:t', namespaces) if node.text]
            return "\n".join(texts)
        except Exception as e:
            logger.warning(f"DOCX XML 고속 추출 실패, DocxDocument 폴백 진행: {e}")
            try:
                doc = DocxDocument(io.BytesIO(file_bytes))
                return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except Exception as e2:
                logger.error(f"DOCX 텍스트 추출 최종 실패: {e2}")
                return ""

    # TXT / MD / CSV 등 텍스트 계열
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="ignore")


def _chunk_text(text: str) -> List[str]:
    """텍스트를 고정 크기 청크로 분할합니다 (CHUNK_OVERLAP 만큼 겹침).

    Args:
        text (str): 원본 텍스트.

    Returns:
        List[str]: 청크 문자열 리스트.
    """
    if not text.strip():
        return []

    chunks: List[str] = []
    start = 0
    while start < len(text):
        end = min(start + CHUNK_SIZE, len(text))
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = end - CHUNK_OVERLAP

    return chunks


class GemFileRag:
    """Gem별 업로드 파일의 RAG 파이프라인을 관리합니다.

    pgvector 백엔드 상에 개별 컬렉션을 구축하고, 문서 업로드 시
    텍스트 추출 및 청킹 후 벡터 데이터베이스에 적재하며 대화 시
    유사 청크를 조회하여 에이전트 답변의 컨텍스트로 공급합니다.
    """

    def __init__(self) -> None:
        """GemFileRag 인스턴스를 초기화합니다."""
        self._embeddings = None

    def get_embeddings(self):
        """임베딩 모델 인스턴스를 지연 초기화로 반환합니다.

        Returns:
            Embeddings: 초기화된 임베딩 생성 모델 인스턴스.
        """
        if self._embeddings is None:
            self._embeddings = init_embeddings(model=EMBED_MODEL)
        return self._embeddings

    def _get_vectorstore(self, gem_id: str) -> PGVector:
        """지정된 Gem ID에 대칭되는 pgvector 스토어 객체를 초기화 및 반환합니다.

        Args:
            gem_id (str): 대상 Gem의 고유 식별자 UUID.

        Returns:
            PGVector: 비동기 모드로 구성된 pgvector 벡터 스토어 연결 인스턴스.
        """
        return PGVector(
            embeddings=self.get_embeddings(),
            collection_name=_collection_name(gem_id),
            connection=CONNECTION,
            async_mode=True,
        )

    async def process_and_store(
        self, gem_id: str, filename: str, file_bytes: bytes
    ) -> int:
        """파일에서 텍스트를 추출하고 청킹 후 pgvector에 저장합니다.

        Args:
            gem_id (str): 대상 Gem의 고유 ID.
            filename (str): 업로드된 파일명.
            file_bytes (bytes): 파일 바이너리 데이터.

        Returns:
            int: 저장된 청크 수. 텍스트를 추출할 수 없으면 0.
        """
        text = _extract_text(filename, file_bytes)
        chunks = _chunk_text(text)

        if not chunks:
            logger.warning(f"파일 '{filename}'에서 텍스트를 추출하지 못했습니다.")
            return 0

        vectorstore = self._get_vectorstore(gem_id)
        metadatas = [{"filename": filename, "chunk_index": i} for i in range(len(chunks))]
        await vectorstore.aadd_texts(chunks, metadatas=metadatas)

        logger.info(f"Gem '{gem_id}' | 파일 '{filename}' → {len(chunks)}개 청크 저장 완료")
        return len(chunks)

    async def search(self, gem_id: str, query: str, k: int = 3) -> List[Dict[str, Any]]:
        """Gem의 업로드 파일 컬렉션에서 유사 청크를 검색합니다.

        Args:
            gem_id (str): 검색할 Gem ID.
            query (str): 검색 쿼리 텍스트.
            k (int): 반환할 최대 청크 수.

        Returns:
            List[Dict[str, Any]]: 각 청크의 filename, chunk_index, text_chunk, score를 담은 딕셔너리 리스트.
        """
        vectorstore = self._get_vectorstore(gem_id)
        results = await vectorstore.asimilarity_search_with_score(query, k=k)

        return [
            {
                "filename": (doc.metadata or {}).get("filename", ""),
                "chunk_index": (doc.metadata or {}).get("chunk_index", 0),
                "text_chunk": doc.page_content,
                "score": round(1.0 - score, 4),
            }
            for doc, score in results
        ]

    async def delete_collection(self, gem_id: str) -> None:
        """Gem의 파일 임베딩 컬렉션 전체를 삭제합니다. Gem 삭제 시 호출합니다.

        Args:
            gem_id (str): 삭제할 Gem ID.
        """
        try:
            vectorstore = self._get_vectorstore(gem_id)
            await vectorstore.adelete_collection()
            logger.info(f"Gem '{gem_id}' 파일 컬렉션 삭제 완료")
        except Exception as exc:
            logger.warning(f"Gem '{gem_id}' 파일 컬렉션 삭제 실패 (무시): {exc}")


# 싱글톤 인스턴스
gem_file_rag = GemFileRag()



